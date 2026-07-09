# Copyright (c) 2026 Baynham Makusha. All rights reserved.
# Unauthorized copying, distribution, or use is prohibited.
"""Resume → MasterProfile extraction.

Shared by the CLI (``cli.import_resume``) and the web onboarding API
(``api.routes.profile``). One Gemini call turns raw resume text into a fully
validated :class:`MasterProfile`; the readers below accept PDF, DOCX, or plain
text so the same path serves a file upload or a pasted resume.
"""

from __future__ import annotations

import io
from pathlib import Path

from google import genai
from google.genai import types

from models.profile import MasterProfile
from obs.llm_cost import record_llm_call
from obs.logging import get_logger

log = get_logger("tools.profile.extract")

# Thinking is "high" by default on Gemini 3.x and bills as output tokens at
# the full output rate. This task tags bullets and infers a few classifications
# but is fundamentally extraction, not open-ended reasoning, so trim rather
# than disable — check obs/llm_cost.py telemetry post-deploy to confirm
# thinking_tokens actually dropped before going lower.
_EXTRACT_THINKING = types.ThinkingConfig(thinking_level=types.ThinkingLevel.LOW)

# Ceiling on what one extraction call may generate — thinking counts toward
# max_output_tokens, so this must cover answer + thinking. Telemetry worst so
# far is ~4.2K combined (1,529 answer, 2,683 thinking) on a normal resume;
# output scales with resume length, so 8192 gives ~2x headroom while capping a
# runaway generation at ~$0.10 instead of the model's ~64K default (~$0.79).
# Hitting the cap truncates the JSON, which fails MasterProfile validation and
# surfaces as extract.gemini.failed rather than a silent partial profile.
_EXTRACT_MAX_OUTPUT_TOKENS = 8192

SYSTEM_PROMPT = """You extract structured career data from resumes.

Rules:
- Preserve the exact wording of bullets. Do not paraphrase or 'improve' them.
- For each bullet, identify 2-5 tags capturing BOTH the technical/domain content AND the
  transferable dimension (leadership, cross-functional work, strategy, 0-to-1 building, etc.).
  Tag the transferable skill even on a technical bullet, so it stays discoverable when
  matching adjacent role families. Use lowercase hyphenated slugs (e.g. 'distributed-systems',
  'cross-functional', 'go-to-market', 'stakeholder-management', 'a-b-testing').
- For each bullet, extract impact only if a number or measurable outcome is stated.
- If a date is just a year, use January 1 of that year for start, December 31 for end.
- For the objective_template, generate a 2-sentence template the user can customize
  per-role, with literal {role} and {company} placeholders. Keep it family-neutral.
- If the resume has no skills section, infer skill categories from the experience bullets.
"""


def read_pdf_text(data: bytes) -> str:
    """Extract text from PDF bytes."""
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_docx_text(data: bytes) -> str:
    """Extract text from DOCX bytes (paragraphs + table cells)."""
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def read_resume_text(data: bytes, filename: str) -> str:
    """Extract plain text from an uploaded resume by sniffing the extension.

    PDF and DOCX are parsed; anything else is decoded as UTF-8 text (covers
    .txt and the "paste resume text" path).
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return read_pdf_text(data)
    if suffix in (".docx", ".doc"):
        return read_docx_text(data)
    return data.decode("utf-8", errors="replace")


def extract_profile(raw_text: str, user_id: str) -> MasterProfile:
    """Run Gemini over resume text and return a validated MasterProfile.

    Deterministic (temperature 0.1) and schema-constrained, so the response is
    a fully-formed :class:`MasterProfile`. ``user_id`` is injected afterward
    because the LLM cannot know it.
    """
    if not raw_text.strip():
        raise ValueError("Resume text is empty — nothing to extract.")

    call_log = log.bind(user_id=user_id, chars=len(raw_text))
    call_log.info("extract.gemini.start", model="gemini-3.1-pro-preview")
    client = genai.Client(vertexai=True)

    try:
        response = client.models.generate_content(
            # "gemini-3-pro" is not a valid catalog id; the Gemini 3 Pro model is
            # "gemini-3.1-pro-preview".
            model="gemini-3.1-pro-preview",
            contents=[raw_text],
            config=types.GenerateContentConfig(
                system_instruction=SYSTEM_PROMPT,
                response_mime_type="application/json",
                response_schema=MasterProfile,
                temperature=0.1,  # we want determinism here
                thinking_config=_EXTRACT_THINKING,
                max_output_tokens=_EXTRACT_MAX_OUTPUT_TOKENS,
            ),
        )
        record_llm_call(step="profile.extract", response=response)
        profile = MasterProfile.model_validate_json(response.text)
    except Exception:
        call_log.exception("extract.gemini.failed")
        raise

    profile.user_id = user_id
    call_log.info("extract.gemini.done", roles=len(profile.experience))
    return profile
