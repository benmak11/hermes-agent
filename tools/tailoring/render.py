# Copyright (c) 2026 Baynham Makusha. All rights reserved.
# Unauthorized copying, distribution, or use is prohibited.
"""Render a tailored resume to .docx and upload it to Cloud Storage.

The document is built programmatically with python-docx using simple paragraph
styles only — no tables or fancy formatting, which ATS parsers (Greenhouse,
Lever, Workday) choke on. Rendering is pure/deterministic; the objective text and
reranked experience are produced upstream by the tailoring pipeline.
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document

from models.profile import Experience, MasterProfile


def _fmt_dates(exp: Experience) -> str:
    start = exp.start.strftime("%b %Y")
    end = exp.end.strftime("%b %Y") if exp.end else "Present"
    return f"{start} - {end}"


def build_resume_doc(
    profile: MasterProfile,
    experience: list[Experience],
    objective: str,
) -> Document:
    """Build (but do not save) the tailored resume document."""
    doc = Document()

    doc.add_heading(profile.full_name, level=0)

    # Contact line — plain text, comma-separated, ATS-friendly.
    contact_bits = [profile.email, profile.phone, profile.location]
    contact_bits += [v for v in profile.links.values() if v]
    contact = " | ".join(b for b in contact_bits if b)
    if contact:
        doc.add_paragraph(contact)

    if objective:
        doc.add_heading("Objective", level=1)
        doc.add_paragraph(objective)

    doc.add_heading("Experience", level=1)
    for exp in experience:
        doc.add_heading(f"{exp.role} — {exp.company}", level=2)
        meta = _fmt_dates(exp)
        if exp.location:
            meta = f"{meta} | {exp.location}"
        doc.add_paragraph(meta)
        for bullet in exp.bullets:
            doc.add_paragraph(bullet.text, style="List Bullet")

    if profile.education:
        doc.add_heading("Education", level=1)
        for edu in profile.education:
            years = (
                f"{edu.start_year}-{edu.end_year}" if edu.end_year else f"{edu.start_year}"
            )
            doc.add_paragraph(f"{edu.degree}, {edu.field} — {edu.institution} ({years})")

    if profile.skills:
        doc.add_heading("Skills", level=1)
        for category, items in profile.skills.items():
            if items:
                doc.add_paragraph(f"{category.title()}: {', '.join(items)}")

    return doc


def render_resume_docx(
    profile: MasterProfile,
    experience: list[Experience],
    objective: str,
    output_path: Path,
) -> Path:
    """Render the tailored resume to a local .docx file and return its path."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build_resume_doc(profile, experience, objective)
    doc.save(str(output_path))
    return output_path


def resume_bucket_name() -> str:
    """Resolve the resumes bucket name (env override or <project>-resumes)."""
    explicit = os.environ.get("RESUME_BUCKET")
    if explicit:
        return explicit
    project = os.environ.get("GOOGLE_CLOUD_PROJECT")
    if not project:
        raise RuntimeError(
            "Set RESUME_BUCKET or GOOGLE_CLOUD_PROJECT to resolve the resumes bucket."
        )
    return f"{project}-resumes"


def upload_resume(local_path: Path, user_id: str, job_id: str) -> str:
    """Upload the rendered resume to GCS and return its gs:// URI."""
    from google.cloud import storage

    client = storage.Client()
    bucket = client.bucket(resume_bucket_name())
    blob = bucket.blob(f"users/{user_id}/applications/{job_id}/resume.docx")
    blob.upload_from_filename(str(local_path))
    return f"gs://{bucket.name}/{blob.name}"
